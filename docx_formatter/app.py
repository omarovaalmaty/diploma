from flask import Flask, render_template, request, send_file
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

# Подсчет метрик
def calculate_page_count(doc):
    page_count = len(doc.element.xpath('//w:sectPr'))
    return page_count

def count_figures(doc):
    figures = len(doc.inline_shapes)
    return figures

def count_tables(doc):
    return len(doc.tables)

def count_sources(doc):
    references_section_started = False
    sources_count = 0
    
    for para in doc.paragraphs:
        if "СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ" in para.text.upper() or \
           "ПАЙДАЛАНЫЛҒАН ДЕРЕККӨЗДЕРДІҢ ТІЗІМІ" in para.text.upper() or \
           "ПАЙДАЛАНЫЛҒАН ӘДЕБИЕТТЕР" in para.text.upper():
            references_section_started = True
            continue
        if references_section_started:
            if re.match(r"^\d+\.", para.text.strip()):
                sources_count += 1
    return sources_count

def count_appendices(doc):
    appendices = 0
    for paragraph in doc.paragraphs:
        if "ПРИЛОЖЕНИЕ" in paragraph.text.upper():
            appendices += 1
    return appendices

def ensure_style(doc, style_name, font_name="Times New Roman", font_size=Pt(12)):
    styles = doc.styles
    if style_name not in styles:
        style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = font_name
        style.font.size = font_size
    return styles[style_name]

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = os.path.join(os.getcwd(), 'uploads')
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # Ограничение размера файла в 16 МБ

# Убедитесь, что папка для загрузок существует
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

def add_referat_page(doc):
    para = doc.paragraphs[0]
    run = para.add_run()
    run.add_break(WD_BREAK.PAGE)  # Добавляем разрыв страницы

    referat_style = ensure_style(doc, 'Heading1', font_size=Pt(14))
    referat_para = doc.add_paragraph("РЕФЕРАТ", style=referat_style)
    referat_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

def format_referat_content(doc, page_count, figure_count, table_count, source_count, appendix_count):
    referat_found = False
    for para in doc.paragraphs:
        if "РЕФЕРАТ" in para.text.upper():
            referat_found = True
            break

    if not referat_found:
        add_referat_page(doc)

    referat_text = f"Отчет состоит из {page_count} страниц, {figure_count} рисунков, {table_count} таблиц, {source_count} источников, {appendix_count} приложений."
    if referat_found:
        para.add_run('\n' + referat_text)
    else:
        doc.add_paragraph(referat_text)

def set_page_numbering(doc):
    sections = doc.sections
    for i, section in enumerate(sections):
        if i == 0:
            continue
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:t')
        fldChar3.text = str(i + 1)
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)
        run._r.append(fldChar4)
        
        if i == 1:
            sectPr = section._sectPr
            pgNumType = OxmlElement('w:pgNumType')
            pgNumType.set(qn('w:start'), "2")
            sectPr.append(pgNumType)

def check_and_format_referat(filepath):
    doc = Document(filepath)
    page_count = calculate_page_count(doc)
    figure_count = count_figures(doc)
    table_count = count_tables(doc)
    source_count = count_sources(doc)
    appendix_count = count_appendices(doc)

    format_referat_content(doc, page_count, figure_count, table_count, source_count, appendix_count)
    
    set_page_numbering(doc)

    formatted_filename = 'formatted_' + os.path.basename(filepath)
    formatted_filepath = os.path.join(app.config['UPLOAD_FOLDER'], formatted_filename)
    doc.save(formatted_filepath)
    return formatted_filepath

@app.route('/', methods=['GET', 'POST'])
def upload():
    if request.method == 'POST':
        file = request.files['file']
        if file:
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            new_filepath = check_and_format_referat(filepath)
            
            return send_file(new_filepath, as_attachment=True)
    
    return render_template('upload.html')

if __name__ == '__main__':
    app.run(debug=True)
